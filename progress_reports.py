import os
import sys
from shutil import copyfile
from docx import Document

if __name__ == '__main__':
    # set path
    path = sys.argv[1]
    os.chdir(path)

    # open roster
    roster = open('roster.txt')

    # open improvements
    improvements = open('improvements.txt')

    # for every student on the roster
    for student in roster:
        # split string
        split_student = student.rstrip().split(" ")

        # set first name
        first_name = split_student[1]

        # set last name
        last_name = split_student[2]

        # set improvement
        for improvement in improvements:
            if improvement.startswith(str.lower(split_student[3])):
                student_improvement = improvement.rstrip()

        # set gender
        gender = split_student[0]

        # set midterm score (when applicable)
        if len(split_student) == 5:
            test_score = split_student[4]

        # make copy of template, rename
        report_filename = last_name + "." + first_name + "_" + sys.argv[2] + ".docx"
        copyfile("Template_" + sys.argv[2] + ".docx", report_filename)
        report = Document(report_filename)

        # for word in template
        for paragraph in report.paragraphs:
            for run in paragraph.runs:
                if not run.bold and not run.italic and not run.underline:
                    # replace first name
                    paragraph.text = paragraph.text.replace('<firstname>', first_name)

                    # replace last name
                    paragraph.text = paragraph.text.replace('<lastname>', last_name)

                    # replace improvement
                    paragraph.text = paragraph.text.replace('<improvement>', student_improvement)

                    # replace midterm score
                    if len(split_student) == 5:
                        paragraph.text = paragraph.text.replace('<testscore>', test_score)

                    # replace pronouns
                    if str.upper(gender) == 'M':
                        paragraph.text = paragraph.text.replace('<spn>', 'he')
                        paragraph.text = paragraph.text.replace('<opn>', 'him')
                        paragraph.text = paragraph.text.replace('<ppn>', 'his')
                    elif str.upper(gender) == 'F':
                        paragraph.text = paragraph.text.replace('<spn>', 'she')
                        paragraph.text = paragraph.text.replace('<opn>', 'her')
                        paragraph.text = paragraph.text.replace('<ppn>', 'hers')
                    else:
                        paragraph.text = paragraph.text.replace('<spn>', 'they')
                        paragraph.text = paragraph.text.replace('<opn>', 'them')
                        paragraph.text = paragraph.text.replace('<ppn>', 'their')

        # fill in student name in table
        report.tables[0].cell(0, 1).paragraphs[0].add_run(first_name + " " + last_name).underline = True

        # save report
        report.save(report_filename)
