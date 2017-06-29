import os
import sys
from shutil import copyfile
from docx import Document

if __name__ == '__main__':
    # set path
    path = sys.argv[1]
    os.chdir(path)

    # open roster
    roster = open('roster.txt')

    for student in roster:
        # split string
        split_student = student.rstrip().split(" ")

        # set first name
        first_name = split_student[1]

        # set last name
        last_name = split_student[2]

        # set gender
        gender = split_student[0]

        # set high school credit
        hsCredit = split_student[4]

        # set percent grade
        pGrade = split_student[3]

        # set letter grade
        if 94 <= float(pGrade) <= 100:
            lGrade = "A"
        elif 90 <= float(pGrade) < 94:
            lGrade = "A-"
        elif 87 <= float(pGrade) < 90:
            lGrade = "B+"
        elif 83 <= float(pGrade) < 87:
            lGrade = "B"
        elif 80 <= float(pGrade) < 83:
            lGrade = "B-"
        elif 77 <= float(pGrade) < 80:
            lGrade = "C+"
        elif 73 <= float(pGrade) < 77:
            lGrade = "C"
        elif 70 <= float(pGrade) < 73:
            lGrade = "C-"
        elif 67 <= float(pGrade) < 70:
            lGrade = "D+"
        elif 63 <= float(pGrade) < 67:
            lGrade = "D"
        elif 60 <= float(pGrade) < 63:
            lGrade = "D-"
        elif 0 <= float(pGrade) < 60:
            lGrade = "F"

        # open comments
        comments = open('comments.txt')

        if float(pGrade) < 80:
            comments.readline()
            comments.readline()
        comment = comments.readline()

        # create new file for student
        copy_filename = last_name + "." + first_name + "_FinalTranscript.docx"  # format filename
        copyfile("Template_FinalTranscript.docx", copy_filename)  # copy file
        report = Document(copy_filename)

        for paragraph in report.paragraphs:
            for run in paragraph.runs:
                if not run.bold and not run.italic and not run.underline:
                    # set letter and percent grade
                    paragraph.text = paragraph.text.replace('<percentgrade>', pGrade)
                    paragraph.text = paragraph.text.replace('<lettergrade>', lGrade)

                    # set high school credit
                    paragraph.text = paragraph.text.replace('<hscredit>', hsCredit)

                    # insert comment section
                    paragraph.text = paragraph.text.replace('<comment>', comment)

                    # replace first name
                    paragraph.text = paragraph.text.replace('<firstname>', first_name)

                    # replace last name
                    paragraph.text = paragraph.text.replace('<lastname>', last_name)

                    # replace pronouns
                    if str.upper(gender) == 'M':
                        paragraph.text = paragraph.text.replace('<spn>', 'he')
                        paragraph.text = paragraph.text.replace('<opn>', 'him')
                        paragraph.text = paragraph.text.replace('<ppn>', 'his')
                    elif str.upper(gender) == 'F':
                        paragraph.text = paragraph.text.replace('<spn>', 'she')
                        paragraph.text = paragraph.text.replace('<opn>', 'her')
                        paragraph.text = paragraph.text.replace('<ppn>', 'hers')
                    else:
                        paragraph.text = paragraph.text.replace('<spn>', 'they')
                        paragraph.text = paragraph.text.replace('<opn>', 'them')
                        paragraph.text = paragraph.text.replace('<ppn>', 'their')

        # fill in student name in header (table) and underline
        report.tables[0].cell(0, 1).paragraphs[0].add_run(first_name + " " + last_name).underline = True

        # save progress report
        report.save(copy_filename)